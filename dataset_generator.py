# dataset_generator_streamlit_fixed.py
"""
Streamlit Synthetic EEG Dataset Generator 
 - top-1 candidate in report (instead of top-3)
 - conservative post-generation band scaling to better match metadata multipliers
 - inline multiplier editor (edit/save multipliers to multipliers_custom.json)
 - heuristic interpretation uses processed RMS data
"""

import os, math, datetime, tempfile, hashlib, json, itertools
from pathlib import Path

import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import seaborn as sns

# optional components
try:
    from sklearn.decomposition import PCA
    from sklearn.cluster import KMeans
    from sklearn.preprocessing import StandardScaler
    from sklearn.metrics import silhouette_score
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

try:
    import docx
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ----------------------------
# Config / persistence
# ----------------------------
CONFIG_PATH = Path("multipliers_custom.json")

def load_custom_multipliers():
    if not CONFIG_PATH.exists():
        return {}
    try:
        with CONFIG_PATH.open("r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict):
                return data
    except Exception:
        pass
    return {}

def save_custom_multipliers(updates: dict):
    """
    updates: dict containing any of the maps, e.g.
    {"MENTAL_MAP": {"newstate": {...}}, "LONG_ISSUE_MAP": {...}}
    This merges into existing file.
    """
    data = load_custom_multipliers()
    # deep merge at top-level keys
    for key, val in updates.items():
        if not isinstance(val, dict):
            continue
        if key not in data or not isinstance(data[key], dict):
            data[key] = {}
        # merge per-entry
        for subk, subv in val.items():
            data[key][subk] = subv
    try:
        with CONFIG_PATH.open("w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        st.warning(f"Could not save custom multipliers: {e}")

def safe_rerun():
    """
    Robust rerun helper: call st.experimental_rerun if available,
    otherwise toggle a session flag and call st.stop() to force UI refresh.
    """
    try:
        rerun = getattr(st, "experimental_rerun", None)
        if callable(rerun):
            rerun()
            return
    except Exception:
        pass
    st.session_state["_rerun_toggle"] = not st.session_state.get("_rerun_toggle", False)
    st.stop()

# ----------------------------
# Channel + band definitions
# ----------------------------
CHANNEL_NAMES = [
    "Fp1","Fp2","F3","F4","F7","F8","Fz",
    "C3","C4","Cz","T3","T4","T5","T6",
    "P3","P4","Pz","O1","O2","Oz",
    "AF3","AF4","FC5","FC6"
]

BAND_ORDER = ["delta","theta","alpha","beta","gamma"]

# (f_lo, f_hi, base_amp_lo, base_amp_hi)
BAND_CFG = {
    "delta": (0.5, 4, 8.0, 40.0),
    "theta": (4, 8, 4.0, 25.0),
    "alpha": (8, 13, 15.0, 55.0),
    "beta":  (13, 30, 4.0, 18.0),
    "gamma": (30, 80, 0.8, 8.0)
}

BASE_ARTIFACT = {"artifact_prob_per_sec": 0.02, "eye_blink_amp_uV": 80.0, "muscle_burst_amp_uV": 30.0}

# ----------------------------
# Literature-informed multipliers (defaults)
# ----------------------------
AGE_GROUP_MAP = {
  "child": {"delta": 1.40, "theta": 1.30, "alpha": 0.75, "beta": 0.85, "gamma": 0.80},
  "young": {"delta": 0.85, "theta": 0.90, "alpha": 1.10, "beta": 1.05, "gamma": 1.10},
  "adult": {"delta": 1.00, "theta": 1.00, "alpha": 1.00, "beta": 1.00, "gamma": 1.00},
  "senior": {"delta": 1.20, "theta": 1.25, "alpha": 0.80, "beta": 0.90, "gamma": 0.70}
}

GENDER_MAP = {
  "male": {"delta": 1.05, "theta": 1.02, "alpha": 0.97, "beta": 1.00, "gamma": 1.08},
  "female": {"delta": 0.95, "theta": 1.00, "alpha": 1.08, "beta": 1.05, "gamma": 0.95},
  "other": {}
}

MENTAL_MAP = {
  "relaxed": {"delta": 0.95, "theta": 1.05, "alpha": 1.30, "beta": 0.85, "gamma": 1.05},
  "neutral": {},
  "stressed": {"delta": 1.00, "theta": 1.10, "alpha": 0.80, "beta": 1.40, "gamma": 1.10},
  "anxious": {"delta": 1.00, "theta": 1.15, "alpha": 0.85, "beta": 1.25, "gamma": 1.05},
  "depressed": {"delta": 1.05, "theta": 1.25, "alpha": 0.80, "beta": 0.90, "gamma": 0.75},
  "happy": {"delta": 0.95, "theta": 1.00, "alpha": 1.15, "beta": 1.10, "gamma": 1.20},
  "drowsy": {"delta": 1.30, "theta": 1.40, "alpha": 0.70, "beta": 0.80, "gamma": 0.70}
}

LONG_ISSUE_MAP = {
    "epilepsy": {"delta": 1.65, "theta": 1.45, "alpha": 0.78, "beta": 0.82, "gamma": 0.88},
    "diabetes": {"delta": 1.15, "theta": 1.20, "alpha": 0.92, "beta": 1.05, "gamma": 0.98},
    "hypertension": {"delta": 1.00, "theta": 1.05, "alpha": 0.88, "beta": 1.22, "gamma": 1.15},
    "heart_disease": {"delta": 1.18, "theta": 1.15, "alpha": 0.90, "beta": 1.08, "gamma": 1.00},
    "post_heart_surgery": {"delta": 1.22, "theta": 1.25, "alpha": 0.93, "beta": 0.92, "gamma": 0.85},
    "asthma": {"delta": 1.05, "theta": 1.18, "alpha": 0.94, "beta": 1.12, "gamma": 1.02},
    "chronic_kidney_disease": {"delta": 1.28, "theta": 1.30, "alpha": 0.88, "beta": 0.92, "gamma": 0.87},
    "thyroid_disorder": {"delta": 1.08, "theta": 1.15, "alpha": 0.93, "beta": 1.10, "gamma": 1.03},
    "arthritis": {"delta": 1.05, "theta": 1.10, "alpha": 0.92, "beta": 1.10, "gamma": 1.00},
    "chronic_liver_disease": {"delta": 1.30, "theta": 1.35, "alpha": 0.85, "beta": 0.90, "gamma": 0.82},
    "depression": {"delta": 1.12, "theta": 1.28, "alpha": 0.85, "beta": 1.18, "gamma": 1.10},
    "anxiety": {"delta": 1.00, "theta": 1.08, "alpha": 0.90, "beta": 1.30, "gamma": 1.22}
}

SHORT_ISSUE_MAP = {
    "fatigue": {"delta": 1.15, "theta": 1.30, "alpha": 1.12, "beta": 0.90, "gamma": 0.88},
    "headache": {"delta": 1.12, "theta": 1.22, "alpha": 0.88, "beta": 1.08, "gamma": 1.02},
    "fracture": {"delta": 1.05, "theta": 1.08, "alpha": 0.95, "beta": 1.10, "gamma": 1.05},
    "fever": {"delta": 1.18, "theta": 1.25, "alpha": 0.92, "beta": 0.95, "gamma": 0.88},
    "typhoid": {"delta": 1.25, "theta": 1.32, "alpha": 0.88, "beta": 0.92, "gamma": 0.85},
    "viral_infection": {"delta": 1.20, "theta": 1.25, "alpha": 0.92, "beta": 1.00, "gamma": 0.95},
    "bacterial_infection": {"delta": 1.22, "theta": 1.28, "alpha": 0.90, "beta": 0.98, "gamma": 0.93},
    "post_surgery_recovery": {"delta": 1.25, "theta": 1.28, "alpha": 0.94, "beta": 0.90, "gamma": 0.85},
    "dehydration": {"delta": 1.15, "theta": 1.18, "alpha": 0.93, "beta": 0.92, "gamma": 0.88},
    "sleep_deprivation": {"delta": 1.30, "theta": 1.40, "alpha": 0.88, "beta": 0.88, "gamma": 0.82},
    "acute_stress": {"delta": 1.00, "theta": 1.05, "alpha": 0.88, "beta": 1.35, "gamma": 1.30},
    "migraine": {"delta": 1.20, "theta": 1.30, "alpha": 0.85, "beta": 1.10, "gamma": 1.05}
}

MUSIC_GENRE_MAP = {
    "classical":{"alpha":1.25,"theta":1.15},
    "jazz":{"alpha":1.12,"theta":1.05},
    "pop":{"beta":1.2},
    "rock":{"beta":1.3,"gamma":1.15},
    "electronic":{"gamma":1.25,"beta":1.1},
    "silence":{}, "none":{}
}

MUSIC_TRANSIENTS = {
    "classical":{"alpha":1.3},
    "rock":{"beta":1.25,"gamma":1.15},
    "electronic":{"gamma":1.25},
    "jazz":{"alpha":1.1},
    "pop":{"beta":1.15}
}

ARTIFACT_STATE = {"stressed":{"artifact_prob_mult":1.8,"muscle_amp_mult":1.5}, "relaxed":{"artifact_prob_mult":0.8}}

# Merge custom multipliers from disk (persisted user additions)
_custom = load_custom_multipliers()
for key, mapping in _custom.items():
    try:
        if key in globals() and isinstance(mapping, dict):
            target = globals()[key]
            if isinstance(target, dict):
                for subk, subv in mapping.items():
                    if isinstance(subv, dict):
                        target[subk] = subv
    except Exception:
        pass

# ----------------------------
# small helpers
# ----------------------------
def norm(s): return "" if s is None else str(s).strip().lower()
def choose_age_group(age):
    try:
        a = int(age)
    except:
        return "adult"
    if a < 13: return "child"
    if a < 25: return "young"
    if a < 60: return "adult"
    return "senior"

def deterministic_seed(metadata):
    key = f"{metadata.get('name','')}_{metadata.get('age','')}_{metadata.get('gender','')}_{metadata.get('mental_state','')}_{metadata.get('music_genre','')}"
    h = hashlib.sha256(key.encode('utf-8')).hexdigest()
    return int(h[:16], 16) % (2**31)

def build_multipliers(metadata):
    bands = {b:1.0 for b in BAND_ORDER}
    artifacts = {"prob_mult":1.0, "eye_amp_mult":1.0, "muscle_amp_mult":1.0}

    ag = choose_age_group(metadata.get("age"))
    for k,v in AGE_GROUP_MAP.get(ag,{}).items(): bands[k] *= v

    g = norm(metadata.get("gender"))
    if g.startswith("m"): eff = GENDER_MAP["male"]
    elif g.startswith("f"): eff = GENDER_MAP["female"]
    else: eff = GENDER_MAP["other"]
    for k,v in eff.items(): bands[k] *= v

    ms = norm(metadata.get("mental_state"))
    for key, mapv in MENTAL_MAP.items():
        if key in ms or ms==key:
            for k,v in mapv.items(): bands[k] *= v
            if key in ARTIFACT_STATE:
                art = ARTIFACT_STATE[key]
                artifacts["prob_mult"] *= art.get("artifact_prob_mult", 1.0)
                artifacts["muscle_amp_mult"] = art.get("muscle_amp_mult", 1.0)
            break

    for it in metadata.get("long_term_issues_list", []):
        for k,v in LONG_ISSUE_MAP.get(it.lower(), {}).items(): bands[k] *= v
    for it in metadata.get("short_term_issues_list", []):
        for k,v in SHORT_ISSUE_MAP.get(it.lower(), {}).items(): bands[k] *= v

    wants = metadata.get("wants_to_listen", False)
    genre = norm(metadata.get("music_genre"))
    if wants:
        for k,v in MUSIC_GENRE_MAP.get(genre, {}).items(): bands[k] *= v
    else:
        bands["alpha"] *= 0.98
        bands["beta"] *= 1.02

    return bands, artifacts

# ----------------------------
# signal components
# ----------------------------
def time_vector(duration, srate):
    n = max(1, int(duration * srate))
    return np.arange(n)/float(srate)

def gen_band_component(t, f_lo, f_hi, a_lo, a_hi, rng):
    sig = np.zeros_like(t)
    n_tones = rng.randint(1,4)
    for _ in range(n_tones):
        f = rng.uniform(f_lo, f_hi)
        phase = rng.uniform(0, 2*np.pi)
        amp = rng.uniform(a_lo, a_hi)
        modf = rng.uniform(0.01, 0.6)
        mod = 0.5*(1 + np.sin(2*np.pi*modf*t + rng.uniform(0, 2*np.pi)))
        sig += amp * mod * np.sin(2*np.pi*f*t + phase)
    sig += rng.normal(scale=max(0.02, np.std(sig)*0.03), size=len(t))
    return sig

def insert_artifacts(sig, t, srate, rng, artifacts):
    n = len(sig)
    prob = (BASE_ARTIFACT["artifact_prob_per_sec"] * artifacts["prob_mult"]) / srate
    for i in range(n):
        if rng.rand() < prob:
            if rng.rand() < 0.6:
                width_s = rng.uniform(0.06, 0.25)
                width = max(1, int(width_s * srate))
                x = np.arange(n)
                bump = np.exp(-0.5 * ((x - i)/(width / 2.5))**2)
                sig += bump * BASE_ARTIFACT["eye_blink_amp_uV"] * artifacts.get("eye_amp_mult", 1.0)
            else:
                width_s = rng.uniform(0.03, 0.18)
                width = max(1, int(width_s * srate))
                start = max(0, i - width//2)
                end = min(n, start + width)
                sig[start:end] += rng.normal(scale=BASE_ARTIFACT["muscle_burst_amp_uV"] * artifacts.get("muscle_amp_mult", 1.0), size=end-start)
    return sig

def time_locked_env(t, onset, dur, genre):
    envs = {b: np.ones_like(t) for b in BAND_ORDER}
    if onset is None or dur is None: return envs
    if genre not in MUSIC_TRANSIENTS: return envs
    if len(t) < 2: return envs
    dt = t[1]-t[0]
    start = max(0, int(onset/dt))
    end = min(len(t), int((onset + dur)/dt))
    if end <= start: return envs
    ramp = max(1, int(0.1*(end-start)))
    up = 0.5 - 0.5*np.cos(np.linspace(0, math.pi, ramp))
    down = up[::-1]
    sustain = np.ones(max(1,(end-start) - 2*ramp))
    win = np.concatenate([up, sustain, down]) if sustain.size>0 else up[:(end-start)]
    for band, mult in MUSIC_TRANSIENTS[genre].items():
        v = np.ones_like(t)
        v[start:end] = 1 + (mult - 1.0) * win
        envs[band] = v
    return envs

# ----------------------------
# inference helpers (updated to top-1)
# ----------------------------
def _vec_norm(v):
    v = np.array(v, dtype=float)
    n = np.linalg.norm(v)
    return v / (n + 1e-12)

def infer_combinations_for_vector_top1(rms_vec, restrict_to_selected=None):
    """
    Return the single best candidate (mental,long,short) with score.
    """
    mental_keys = list(MENTAL_MAP.keys())
    long_keys = list(LONG_ISSUE_MAP.keys()) + ["none"]
    short_keys = list(SHORT_ISSUE_MAP.keys()) + ["none"]

    # limit mental key search for speed
    max_ment = min(12, len(mental_keys))
    mental_keys = mental_keys[:max_ment]

    obs_n = _vec_norm(np.array(rms_vec, dtype=float))
    best_score = -2.0
    best_trip = (None, None, None)
    for m in mental_keys:
        m_map = MENTAL_MAP.get(m, {})
        for l in long_keys:
            l_map = LONG_ISSUE_MAP.get(l, {}) if l!="none" else {}
            for s in short_keys:
                s_map = SHORT_ISSUE_MAP.get(s, {}) if s!="none" else {}
                comb = []
                for b in BAND_ORDER:
                    comb.append(m_map.get(b, 1.0) * l_map.get(b, 1.0) * s_map.get(b, 1.0))
                comb_n = _vec_norm(comb)
                sim = float(np.dot(obs_n, comb_n))
                boost = 0.0
                if restrict_to_selected:
                    if restrict_to_selected.get("mental") and restrict_to_selected["mental"]==m:
                        boost += 0.04  # slight increase to favour selected metadata
                    if restrict_to_selected.get("long") and restrict_to_selected["long"]==l:
                        boost += 0.02
                    if restrict_to_selected.get("short") and restrict_to_selected["short"]==s:
                        boost += 0.02
                score = sim + boost
                if score > best_score:
                    best_score = score
                    best_trip = (m, l, s)
    # format output
    m,l,s = best_trip
    combo = []
    if m: combo.append(m)
    if l and l!="none": combo.append(l)
    if s and s!="none": combo.append(s)
    label = ", ".join(combo) if combo else "none"
    return f"{label} ({best_score:.3f})"

def heuristic_interpretation_from_rms(rms_df):
    """
    Build an overall interpretation string from processed RMS data (per-band averages).
    Also returns per-channel short notes dictionary.
    """
    mean_band = rms_df.mean(axis=0)
    overall = []
    dom = mean_band.idxmax()
    if dom == "alpha":
        overall.append("Alpha dominant overall: participant shows relaxed/eyes-closed style activity.")
    elif dom == "beta":
        overall.append("Beta dominant overall: increased alertness, cognitive load, or EMG contamination possible.")
    elif dom in ("theta","delta"):
        overall.append("Slow-wave dominance overall: drowsiness or slow-wave activity; interpret cautiously.")
    elif dom == "gamma":
        overall.append("Gamma elevated overall: higher cognitive engagement or positive affect.")
    else:
        overall.append("Mixed band activity overall.")

    # per-channel quick notes (dominant band per channel)
    per_channel_notes = {}
    for ch in rms_df.index:
        chseries = rms_df.loc[ch]
        cdom = chseries.idxmax()
        note = ""
        if cdom == "alpha":
            note = "alpha-dominant (relaxed/eyes-closed)"
        elif cdom == "beta":
            note = "beta-dominant (alert/muscle-artifact)"
        elif cdom in ("theta","delta"):
            note = "slow-wave dominant (drowsy/slow-wave)"
        elif cdom == "gamma":
            note = "gamma-elevated (active cognition)"
        else:
            note = "mixed"
        per_channel_notes[ch] = note
    return " ".join(overall), per_channel_notes

# ----------------------------
# main dataset generator (with post-scaling to better match metadata multipliers)
# ----------------------------
def generate_dataset(metadata, duration, srate, template_path=None, onset=None, mdur=None, seed=None):
    # deterministic seed if none provided
    if seed is None:
        seed = deterministic_seed(metadata)
    rng = np.random.RandomState(seed)

    t = time_vector(duration, srate)
    n = len(t)

    band_mult, artifacts = build_multipliers(metadata)
    envs = time_locked_env(t, onset, mdur, norm(metadata.get("music_genre","none")))

    cols = ["time"] + [f"{ch}_{b}" for ch in CHANNEL_NAMES for b in BAND_ORDER]
    arr = np.zeros((n, len(cols)))
    arr[:,0] = t

    for ci, ch in enumerate(CHANNEL_NAMES):
        chscale = rng.uniform(0.9, 1.1)
        for bi, band in enumerate(BAND_ORDER):
            f_lo, f_hi, a_lo, a_hi = BAND_CFG[band]
            mult = chscale * band_mult.get(band, 1.0)
            a_lo_adj = a_lo * mult
            a_hi_adj = a_hi * mult
            sig = gen_band_component(t, f_lo, f_hi, a_lo_adj, a_hi_adj, rng)
            sig = sig * envs.get(band, np.ones_like(t))
            col = 1 + ci*len(BAND_ORDER) + bi
            arr[:, col] = sig

    # channel-level sum -> insert artifacts on sum then redistribute
    chan_sum = arr[:,1:].reshape(n, len(CHANNEL_NAMES), len(BAND_ORDER)).sum(axis=2)
    for ci in range(len(CHANNEL_NAMES)):
        chan_sum[:,ci] = insert_artifacts(chan_sum[:,ci], t, srate, rng, artifacts)

    for ci in range(len(CHANNEL_NAMES)):
        idxs = [1 + ci*len(BAND_ORDER) + b for b in range(len(BAND_ORDER))]
        original = arr[:, idxs]
        denom = np.abs(original).sum(axis=1, keepdims=True) + 1e-12
        props = np.abs(original) / denom
        arr[:, idxs] = props * chan_sum[:,ci][:, None] * np.sign(original + 1e-12)

    df = pd.DataFrame(arr, columns=cols)

    # ---- post-generation conservative band scaling to nudge RMS to match metadata multipliers ----
    # compute current RMS per band (averaged across channels)
    rms_map = {}
    for ch in CHANNEL_NAMES:
        rms_map[ch] = {}
        for b in BAND_ORDER:
            col = f"{ch}_{b}"
            rms_map[ch][b] = float(np.sqrt(np.mean(np.square(df[col].values))))

    rms_df = pd.DataFrame({ch: rms_map[ch] for ch in CHANNEL_NAMES}).T
    mean_band_obs = rms_df.mean(axis=0)  # per-band observed mean

    # target relative strengths from band_mult (normalize)
    target = np.array([band_mult.get(b, 1.0) for b in BAND_ORDER], dtype=float)
    if np.all(target > 0):
        target_rel = target / (target.sum() + 1e-12)
        obs_rel = mean_band_obs.values / (mean_band_obs.values.sum() + 1e-12)
        # compute band-wise scaling factors to align obs_rel toward target_rel
        # scale = (target_rel/obs_rel) ** 0.5 (sqrt to convert energy->amplitude approx)
        raw_scale = np.divide(target_rel, obs_rel + 1e-12)
        amp_scale = np.sqrt(raw_scale)
        # clamp scale to avoid extreme changes (e.g., between 0.7 and 1.4)
        amp_scale = np.clip(amp_scale, 0.7, 1.4)
        # apply scaling per band across all channels (multiply columns)
        for bi, b in enumerate(BAND_ORDER):
            s = float(amp_scale[bi])
            # multiply all channel-band columns by s
            for ci, ch in enumerate(CHANNEL_NAMES):
                col = f"{ch}_{b}"
                df[col] = df[col] * s

        # recompute rms_map/rms_df after scaling (for reporting)
        for ch in CHANNEL_NAMES:
            for b in BAND_ORDER:
                col = f"{ch}_{b}"
                rms_map[ch][b] = float(np.sqrt(np.mean(np.square(df[col].values))))
        rms_df = pd.DataFrame({ch: rms_map[ch] for ch in CHANNEL_NAMES}).T

    # append metadata columns at end (ensures report shows what was used)
    df["patient_name"] = metadata.get("name","")
    df["age"] = metadata.get("age","")
    df["gender"] = metadata.get("gender","")
    df["mental_state"] = metadata.get("mental_state","")
    df["long_term_issues"] = ", ".join(metadata.get("long_term_issues_list", []))
    df["short_term_issues"] = ", ".join(metadata.get("short_term_issues_list", []))
    df["wants_to_listen"] = metadata.get("wants_to_listen", False)
    df["music_genre"] = metadata.get("music_genre","")
    if onset is not None: df["music_onset"] = onset
    if mdur is not None: df["music_duration"] = mdur

    # template reorder if provided
    if template_path:
        try:
            tmp = pd.read_csv(template_path, nrows=0)
            tmpcols = list(tmp.columns)
            # only reorder if template columns all present
            if all(c in df.columns for c in tmpcols):
                df = df[tmpcols]
        except Exception:
            pass

    name_safe = "".join([c for c in str(metadata.get("name","subject")) if c.isalnum() or c in (" ","_")]).strip().replace(" ","_")
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    fname = f"synthetic_eeg_{name_safe}_{ts}.csv"
    df.to_csv(fname, index=False)
    # return rms_df also for convenience if caller wants it
    return df, fname, rms_df

# ----------------------------
# simple docx report code (uses processed RMS and top-1 combos)
# ----------------------------
def generate_docx_report(df, metadata, output_path):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    # compute RMS per channel-band (if generate_dataset returned rms_df, caller can pass it separately)
    rms_map = {}
    for ch in CHANNEL_NAMES:
        rms_map[ch] = {}
        for b in BAND_ORDER:
            col = f"{ch}_{b}"
            if col in df.columns:
                rms_map[ch][b] = float(np.sqrt(np.mean(np.square(df[col].values))))
            else:
                rms_map[ch][b] = 0.0

    rms_df = pd.DataFrame({ch: rms_map[ch] for ch in CHANNEL_NAMES}).T

    # Build best candidate per channel (top-1)
    combos_per_channel = {}
    restrict = {
        "mental": norm(metadata.get("mental_state","")).lower() or None,
        "long": metadata.get("long_term_issues_list", [None])[0] if metadata.get("long_term_issues_list") else None,
        "short": metadata.get("short_term_issues_list", [None])[0] if metadata.get("short_term_issues_list") else None
    }
    for ch in CHANNEL_NAMES:
        vec = [rms_map[ch][b] for b in BAND_ORDER]
        best = infer_combinations_for_vector_top1(vec, restrict_to_selected=restrict)
        combos_per_channel[ch] = best

    # correlation heatmap
    corr = rms_df.fillna(0).T.corr()
    fig, ax = plt.subplots(figsize=(7,6))
    sns.heatmap(corr, ax=ax, cmap="coolwarm", vmin=-1, vmax=1)
    heat_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.tight_layout()
    fig.savefig(heat_tmp.name, dpi=150)
    plt.close(fig)

    # bar chart average band
    mean_band = rms_df.mean(axis=0)
    fig2, ax2 = plt.subplots(figsize=(6,3))
    mean_band.plot.bar(ax=ax2)
    ax2.set_ylabel("RMS (µV)")
    bar_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig2.tight_layout()
    fig2.savefig(bar_tmp.name, dpi=150)
    plt.close(fig2)

    # interpretations
    overall_text, per_channel_notes = heuristic_interpretation_from_rms(rms_df)

    # create docx
    doc = docx.Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    doc.add_heading(f"{metadata.get('name','Subject')} — EEG Report", level=1)
    doc.add_paragraph("Generated: " + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    tbl = doc.add_table(rows=0, cols=2)
    def add_row(k,v):
        r = tbl.add_row().cells
        r[0].text = str(k); r[1].text = str(v)
    add_row("Name", metadata.get("name",""))
    add_row("Age", metadata.get("age",""))
    add_row("Gender", metadata.get("gender",""))
    add_row("Mental state", metadata.get("mental_state",""))
    add_row("Long-term issues", ", ".join(metadata.get("long_term_issues_list", [])))
    add_row("Short-term issues", ", ".join(metadata.get("short_term_issues_list", [])))
    add_row("Wants to listen", str(metadata.get("wants_to_listen", False)))
    add_row("Music genre", metadata.get("music_genre",""))
    add_row("Music onset (s)", str(metadata.get("music_onset","")))
    add_row("Music duration (s)", str(metadata.get("music_duration","")))

    doc.add_page_break()
    doc.add_heading("Band RMS (per channel) + Best candidate", level=2)
    cols_count = 1 + len(BAND_ORDER) + 1
    table = doc.add_table(rows=1, cols=cols_count, style='Light Grid Accent 1')
    hdr = table.rows[0].cells
    hdr[0].text = "Channel"
    for i,b in enumerate(BAND_ORDER):
        hdr[i+1].text = b.upper()
    hdr[len(BAND_ORDER)+1].text = "BEST_CANDIDATE"

    for ch in CHANNEL_NAMES:
        cells = table.add_row().cells
        cells[0].text = ch
        for i,b in enumerate(BAND_ORDER):
            cells[i+1].text = f"{rms_map[ch][b]:.3f}"
        cells[len(BAND_ORDER)+1].text = combos_per_channel.get(ch, "")

    doc.add_paragraph()
    doc.add_picture(bar_tmp.name, width=Inches(6))
    doc.add_page_break()
    doc.add_heading("Inter-channel correlation (band RMS)", level=2)
    doc.add_paragraph("Heatmap of Pearson correlation between channel band RMS vectors.")
    doc.add_picture(heat_tmp.name, width=Inches(6))
    doc.add_page_break()
    doc.add_heading("Interpretation (heuristic)", level=2)
    doc.add_paragraph(f"Overall: {overall_text}")
    doc.add_paragraph("Per-channel quick notes:", style='List Bullet')
    for ch in CHANNEL_NAMES:
        doc.add_paragraph(f"{ch}: {per_channel_notes.get(ch,'')}", style='List Number')

    doc.add_paragraph()
    doc.add_paragraph("Notes: This interpretation is heuristic and provided for exploratory research. For clinical conclusions, consult domain experts and use clinical-grade recordings and validated pipelines.")

    doc.save(output_path)
    try:
        os.unlink(heat_tmp.name); os.unlink(bar_tmp.name)
    except: pass
    return output_path

# ----------------------------
# Streamlit UI: metadata, add/edit multipliers UI
# ----------------------------
st.set_page_config(page_title="Synthetic EEG Generator (fixed)", layout="wide")
st.title("Synthetic EEG Dataset Generator")

st.sidebar.header("Participant metadata (enter values then press Generate)")

name = st.sidebar.text_input("Name", "")
age = st.sidebar.number_input("Age", min_value=1, max_value=120, value=25)
gender = st.sidebar.selectbox("Gender", ("Male","Female","Other"))

# prepare session-state-backed lists for addable entries
if "mental_states" not in st.session_state:
    st.session_state["mental_states"] = [m for m in MENTAL_MAP.keys()]

# ensure newly added state is auto-selected after rerun
default_ms = st.session_state.get("mental_states_default", st.session_state["mental_states"][0])
if "mental_states_new_selected" in st.session_state:
    if st.session_state["mental_states_new_selected"] in st.session_state["mental_states"]:
        default_ms = st.session_state["mental_states_new_selected"]
    st.session_state.pop("mental_states_new_selected", None)

mental_state = st.sidebar.selectbox("Mental state", st.session_state["mental_states"], index=st.session_state["mental_states"].index(default_ms) if default_ms in st.session_state["mental_states"] else 0)
new_ms = st.sidebar.text_input("Add mental state (type & press Add)", key="add_ms")
if st.sidebar.button("Add mental state"):
    v = new_ms.strip()
    if v:
        key = v.lower()
        if key not in MENTAL_MAP:
            default_bands = {b:1.0 for b in BAND_ORDER}
            MENTAL_MAP[key] = default_bands
            save_custom_multipliers({"MENTAL_MAP": {key: default_bands}})
            st.session_state["mental_states"].append(key)
            st.session_state["mental_states_new_selected"] = key
        else:
            st.warning(f"'{v}' already exists in mental states.")
    else:
        st.warning("Type a mental state name first.")
    safe_rerun()

st.sidebar.markdown("### Long-term issues (multi-select; add new then select)")
if "lt_issues" not in st.session_state:
    st.session_state["lt_issues"] = sorted(list(LONG_ISSUE_MAP.keys()))

default_lt = st.session_state.get("lt_issues_default", st.session_state["lt_issues"][0] if st.session_state["lt_issues"] else None)
if "lt_issues_new_selected" in st.session_state:
    if st.session_state["lt_issues_new_selected"] in st.session_state["lt_issues"]:
        default_lt = st.session_state["lt_issues_new_selected"]
    st.session_state.pop("lt_issues_new_selected", None)

new_lt = st.sidebar.text_input("Add long-term issue", key="new_lt")
if st.sidebar.button("Add long-term issue"):
    v = new_lt.strip()
    if v:
        key = v.lower()
        if key not in LONG_ISSUE_MAP:
            default_bands = {b:1.0 for b in BAND_ORDER}
            LONG_ISSUE_MAP[key] = default_bands
            save_custom_multipliers({"LONG_ISSUE_MAP": {key: default_bands}})
            st.session_state["lt_issues"].append(key)
            st.session_state["lt_issues_new_selected"] = key
        else:
            st.warning(f"'{v}' already exists in long-term issues.")
    else:
        st.warning("Type an issue first.")
    safe_rerun()
long_selected = st.sidebar.multiselect("Long-term issues", options=sorted(st.session_state["lt_issues"]))

st.sidebar.markdown("### Short-term issues (multi-select; add new then select)")
if "st_issues" not in st.session_state:
    st.session_state["st_issues"] = sorted(list(SHORT_ISSUE_MAP.keys()))

default_st = st.session_state.get("st_issues_default", st.session_state["st_issues"][0] if st.session_state["st_issues"] else None)
if "st_issues_new_selected" in st.session_state:
    if st.session_state["st_issues_new_selected"] in st.session_state["st_issues"]:
        default_st = st.session_state["st_issues_new_selected"]
    st.session_state.pop("st_issues_new_selected", None)

new_st = st.sidebar.text_input("Add short-term issue", key="new_st")
if st.sidebar.button("Add short-term issue"):
    v = new_st.strip()
    if v:
        key = v.lower()
        if key not in SHORT_ISSUE_MAP:
            default_bands = {b:1.0 for b in BAND_ORDER}
            SHORT_ISSUE_MAP[key] = default_bands
            save_custom_multipliers({"SHORT_ISSUE_MAP": {key: default_bands}})
            st.session_state["st_issues"].append(key)
            st.session_state["st_issues_new_selected"] = key
        else:
            st.warning(f"'{v}' already exists in short-term issues.")
    else:
        st.warning("Type an issue first.")
    safe_rerun()
short_selected = st.sidebar.multiselect("Short-term issues", options=sorted(st.session_state["st_issues"]))

st.sidebar.markdown("---")
st.sidebar.header("Music options")
wants_to_listen = st.sidebar.selectbox("Wants to listen?", ("True","False")) == "True"
genre_choices = list(MUSIC_GENRE_MAP.keys())
music_genre = st.sidebar.selectbox("Genre", genre_choices, index=genre_choices.index("classical"))
time_of_day = st.sidebar.selectbox("Time of day", ("morning","afternoon","evening","night"))
music_onset = st.sidebar.number_input("Music onset (s)", min_value=0.0, value=10.0)
music_duration = st.sidebar.number_input("Music duration (s)", min_value=0.0, value=30.0)

st.sidebar.markdown("---")
st.sidebar.header("Recording settings")
rec_duration = st.sidebar.number_input("Recording duration (s)", min_value=1.0, value=60.0)
srate = st.sidebar.selectbox("Sample rate (Hz)", (128,256,512), index=1)
template_file = st.sidebar.file_uploader("Template CSV (optional)", type=["csv"])
batch_file = st.sidebar.file_uploader("Batch metadata CSV (optional)", type=["csv"])

st.sidebar.markdown("---")
create_docx = st.sidebar.checkbox("Create .docx report", value=False)
validate_opt = st.sidebar.checkbox("Enable sklearn quick validate (PCA+KMeans)", value=False)

# ----------------------------
# Inline multiplier editor (new)
# ----------------------------
st.sidebar.markdown("---")
st.sidebar.header("Multiplier editor")
with st.sidebar.expander("Open multiplier editor"):
    map_choice = st.selectbox("Map to edit", ("MENTAL_MAP","LONG_ISSUE_MAP","SHORT_ISSUE_MAP"))
    # choose keys
    if map_choice == "MENTAL_MAP":
        keys = sorted(list(MENTAL_MAP.keys()))
    elif map_choice == "LONG_ISSUE_MAP":
        keys = sorted(list(LONG_ISSUE_MAP.keys()))
    else:
        keys = sorted(list(SHORT_ISSUE_MAP.keys()))
    selected_key = st.selectbox("Select entry to edit (or Add new)", ["--add new--"] + keys)
    if selected_key == "--add new--":
        new_entry_name = st.text_input("New key name (type then press Create)", key="create_key_name")
        if st.button("Create new entry"):
            nk = new_entry_name.strip().lower()
            if nk:
                default_bands = {b:1.0 for b in BAND_ORDER}
                if map_choice == "MENTAL_MAP":
                    MENTAL_MAP[nk] = default_bands
                elif map_choice == "LONG_ISSUE_MAP":
                    LONG_ISSUE_MAP[nk] = default_bands
                else:
                    SHORT_ISSUE_MAP[nk] = default_bands
                save_custom_multipliers({map_choice: {nk: default_bands}})
                st.success(f"Created {nk} in {map_choice}")
                safe_rerun()
            else:
                st.warning("Enter a non-empty name.")
    else:
        # display editable numeric inputs for bands for selected_key
        if map_choice == "MENTAL_MAP":
            entry_map = MENTAL_MAP.get(selected_key, {b:1.0 for b in BAND_ORDER})
        elif map_choice == "LONG_ISSUE_MAP":
            entry_map = LONG_ISSUE_MAP.get(selected_key, {b:1.0 for b in BAND_ORDER})
        else:
            entry_map = SHORT_ISSUE_MAP.get(selected_key, {b:1.0 for b in BAND_ORDER})

        st.write(f"Editing: {selected_key} in {map_choice}")
        updated = {}
        cols = st.columns(len(BAND_ORDER))
        for i,b in enumerate(BAND_ORDER):
            with cols[i]:
                val = float(entry_map.get(b, 1.0))
                nv = st.number_input(f"{b}", value=val, key=f"{map_choice}_{selected_key}_{b}", format="%.3f")
                updated[b] = nv
        if st.button("Save changes to JSON"):
            save_custom_multipliers({map_choice: {selected_key: updated}})
            # update in-memory map
            if map_choice == "MENTAL_MAP":
                MENTAL_MAP[selected_key] = updated
            elif map_choice == "LONG_ISSUE_MAP":
                LONG_ISSUE_MAP[selected_key] = updated
            else:
                SHORT_ISSUE_MAP[selected_key] = updated
            st.success("Saved multipliers.")
            safe_rerun()

# ----------------------------
# Main actions
# ----------------------------
st.header("Actions")
col1, col2 = st.columns([1,1])

with col1:
    if st.button("Generate single dataset (use current metadata)"):
        metadata = {
            "name": name or "subject",
            "age": int(age),
            "gender": gender,
            "mental_state": mental_state,
            "long_term_issues_list": long_selected,
            "short_term_issues_list": short_selected,
            "wants_to_listen": wants_to_listen,
            "music_genre": music_genre,
            "time_of_day": time_of_day,
            "music_onset": music_onset,
            "music_duration": music_duration
        }
        # write template to disk if uploaded
        template_path = None
        if template_file is not None:
            tf = tempfile.NamedTemporaryFile(delete=False, suffix=".csv")
            tf.write(template_file.getvalue())
            tf.flush(); tf.close()
            template_path = tf.name
        try:
            # updated generate_dataset returns (df, fname, rms_df)
            df, fname, rms_df = generate_dataset(metadata, rec_duration, srate, template_path, music_onset, music_duration)
            st.session_state["last_df"] = df
            st.session_state["last_fname"] = fname
            st.session_state["last_rms_df"] = rms_df
            st.success(f"Generated: {fname}")
            if create_docx:
                if not DOCX_AVAILABLE:
                    st.error("python-docx not installed; cannot create report.")
                else:
                    docx_file = fname.replace(".csv", ".docx")
                    generate_docx_report(df, metadata, docx_file)
                    st.success(f".docx report created: {docx_file}")
        except Exception as e:
            st.error("Generation failed: " + str(e))

    if st.button("Generate batch (from uploaded batch CSV)"):
        if batch_file is None:
            st.error("Upload a batch CSV in the sidebar first.")
        else:
            try:
                batch_df = pd.read_csv(batch_file)
                outdir = "generated_batch"
                os.makedirs(outdir, exist_ok=True)
                saved = []
                for i, row in batch_df.iterrows():
                    md = {
                        "name": row.get("name", row.get("Name", f"subject_{i}")),
                        "age": int(row.get("age", row.get("Age", 25))),
                        "gender": row.get("gender", row.get("Gender", "Other")),
                        "mental_state": row.get("mental_state", row.get("Mental_State", "neutral")),
                        "long_term_issues_list": [s.strip() for s in str(row.get("long_term_issues", "")).split(",") if s.strip()],
                        "short_term_issues_list": [s.strip() for s in str(row.get("short_term_issues", "")).split(",") if s.strip()],
                        "wants_to_listen": bool(str(row.get("wants_to_listen", False)).strip().lower() in ("1","true","yes")),
                        "music_genre": row.get("music_genre", row.get("Music_genre", "none")),
                        "time_of_day": row.get("time_of_day", row.get("Time_of_day", "evening")),
                        "music_onset": float(row.get("music_onset", 0.0)) if pd.notna(row.get("music_onset", None)) else 0.0,
                        "music_duration": float(row.get("music_duration", 0.0)) if pd.notna(row.get("music_duration", None)) else 0.0
                    }
                    dfg, fname, _ = generate_dataset(md, rec_duration, srate, None, md["music_onset"], md["music_duration"], seed=1000+i)
                    dest = os.path.join(outdir, fname)
                    dfg.to_csv(dest, index=False)
                    saved.append(dest)
                    if create_docx and DOCX_AVAILABLE:
                        docx_name = dest.replace(".csv", ".docx")
                        generate_docx_report(dfg, md, docx_name)
                st.success(f"Batch generated {len(saved)} files in {outdir}")
            except Exception as e:
                st.error("Batch failed: " + str(e))

with col2:
    st.subheader("Preview")
    if "last_df" in st.session_state:
        ldf = st.session_state["last_df"]
        st.write("First 6 rows of last generated file:")
        st.dataframe(ldf.head(6))
        st.download_button("Download CSV", ldf.to_csv(index=False).encode("utf-8"), file_name=st.session_state.get("last_fname","synthetic.csv"), mime="text/csv")
        if "last_rms_df" in st.session_state:
            st.write("Per-channel RMS (summary):")
            st.dataframe(st.session_state["last_rms_df"].round(3))
        if validate_opt and SKLEARN_AVAILABLE:
            if st.button("Quick validate last dataset (PCA+KMeans)"):
                try:
                    n = len(ldf)
                    epoch_samps = max(1, int(srate*2))
                    n_epoch = n // epoch_samps
                    if n_epoch < 2:
                        st.warning("Not enough epochs for validation.")
                    else:
                        feats=[]
                        for e in range(n_epoch):
                            s = e*epoch_samps; e2 = s+epoch_samps
                            vec=[]
                            for ch in CHANNEL_NAMES:
                                for b in BAND_ORDER:
                                    col = f"{ch}_{b}"
                                    seg = ldf[col].values[s:e2] if col in ldf.columns else np.zeros(e2-s)
                                    vec.append(np.sqrt(np.mean(np.square(seg))))
                            feats.append(vec)
                        X = np.array(feats)
                        Xs = StandardScaler().fit_transform(X)
                        pca = PCA(n_components=min(2, Xs.shape[1]))
                        Z = pca.fit_transform(Xs)
                        k = min(3, max(2, n_epoch//2))
                        km = KMeans(n_clusters=k, random_state=0).fit(Z)
                        sil = silhouette_score(Z, km.labels_) if len(set(km.labels_))>1 else -1
                        st.write("PCA explained variance:", pca.explained_variance_ratio_)
                        st.write(f"Silhouette (k={k}): {sil:.3f}")
                        fig, ax = plt.subplots()
                        ax.scatter(Z[:,0], Z[:,1], c=km.labels_, cmap='tab10')
                        st.pyplot(fig)
                except Exception as e:
                    st.error("Validation failed: " + str(e))
        elif validate_opt and not SKLEARN_AVAILABLE:
            st.info("Install scikit-learn to enable validation.")
    else:
        st.info("No dataset generated. Click 'Generate single dataset'.")

st.markdown("---")
# st.markdown("Notes: The generator uses literature-informed multipliers (music -> alpha/theta for relaxing music, anxiety -> beta, depression -> increased theta/reduced alpha, age -> alpha decline). The report now includes a single best candidate per channel (best match to RMS shape). The generator performs a conservative post-synthesis band scaling to nudge generated RMS towards metadata-implied multipliers. This is heuristic/exploratory; consult domain experts for clinical use.")
